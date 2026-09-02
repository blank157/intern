"""Converter + parser-schema tests for the answer-key pipeline."""

from __future__ import annotations

import asyncio
import io
import zipfile

import pytest

from answer_eval.answerkey.converters import (
    UnsupportedAnswerKeyError,
    convert_source,
    detect_format,
)
from answer_eval.answerkey.parser import FakeAnswerKeyParserAgent
from answer_eval.answerkey.schemas import ParsedAnswerKey


def _docx_bytes(paragraphs: list[str]) -> bytes:
    import docx

    container = docx.Document()
    for text in paragraphs:
        container.add_paragraph(text)
    buffer = io.BytesIO()
    container.save(buffer)
    return buffer.getvalue()


def _pdf_bytes(text: str) -> bytes:
    import fitz

    doc = fitz.open()
    doc.new_page().insert_text((72, 72), text)
    data = doc.tobytes()
    doc.close()
    return data


def test_detect_format_rejects_unknown() -> None:
    assert detect_format("key.PDF") == "pdf"
    with pytest.raises(UnsupportedAnswerKeyError):
        detect_format("key.xlsx")


def test_pdf_conversion_extracts_text_per_page() -> None:
    document = convert_source("key.pdf", _pdf_bytes("Q1. Define TCP.\nAnswer: protocol"))
    assert document.format == "pdf"
    assert len(document.pages) == 1
    assert "TCP" in document.pages[0].text
    assert "Q1" in document.full_text()
    # Pages are rendered so scanned/image-only keys reach the VLM.
    assert document.pages[0].image_bytes
    assert document.pages[0].image_bytes.startswith(b"\x89PNG")


def test_docx_conversion_includes_paragraphs_and_tables() -> None:
    import docx

    container = docx.Document()
    container.add_paragraph("1. What is overfitting?")
    table = container.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Concept"
    table.rows[0].cells[1].text = "2 marks"
    buffer = io.BytesIO()
    container.save(buffer)
    document = convert_source("key.docx", buffer.getvalue())
    assert "overfitting" in document.pages[0].text
    assert "Concept | 2 marks" in document.pages[0].text


def test_image_conversion_produces_single_page() -> None:
    from PIL import Image

    image = Image.new("RGB", (80, 40), color=(255, 255, 255))
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    document = convert_source("key.png", buffer.getvalue())
    assert document.format == "png"
    assert len(document.pages) == 1 and document.pages[0].image_bytes


def test_legacy_doc_without_libreoffice_fails_clearly() -> None:
    from answer_eval.answerkey.converters import shutil as conv_shutil  # not exported; use direct call instead

    del conv_shutil
    # Directly exercise converter: environment may or may not have soffice.
    try:
        document = convert_source("key.doc", b"\xd0\xcf\x11\xe0 legacy")
        assert document.format == "doc"  # LibreOffice available in env
    except UnsupportedAnswerKeyError as exc:
        assert "LibreOffice" in str(exc)


def test_zip_disguised_as_docx_rejected() -> None:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("evil.txt", "nope")
    with pytest.raises(UnsupportedAnswerKeyError):
        convert_source("key.docx", buffer.getvalue())


def test_fake_parser_returns_valid_schema() -> None:
    agent = FakeAnswerKeyParserAgent()
    parsed = asyncio.run(agent.parse(convert_source("key.pdf", _pdf_bytes("fixture"))))
    assert isinstance(parsed, ParsedAnswerKey)
    assert parsed.question_count == len(parsed.questions)
    assert parsed.questions[0].expected_concepts[0].concept_code == "C1"


def test_parsed_key_validation_flags_mismatched_totals() -> None:
    from answer_eval.answerkey.schemas import ParsedQuestion

    key = ParsedAnswerKey(
        question_count=5,  # wrong on purpose
        questions=[
            ParsedQuestion(question_number=1, maximum_marks=4),
            ParsedQuestion(question_number=1, maximum_marks=4),  # duplicate number
        ],
    )
    validated = key.validated()
    assert validated.question_count == 2
    joined = " ".join(validated.parser_warnings)
    assert "Duplicate question numbers" in joined
    assert "question_count field said 5" in joined
