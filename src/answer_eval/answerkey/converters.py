"""Convert answer-key source files (PDF/DOC/DOCX/images) into one
representation: ordered pages, each with optional text and an image.

The frontend never needs to know the input format.
"""

from __future__ import annotations

import io
import logging
import shutil
import subprocess  # noqa: S404 - controlled binary invocation (LibreOffice)
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

from PIL import Image

logger = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {".pdf", ".doc", ".docx", ".jpg", ".jpeg", ".png", ".webp"}
_IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp"}


class UnsupportedAnswerKeyError(Exception):
    """Raised when the file cannot be converted at all."""


@dataclass
class SourcePage:
    page_number: int
    text: str = ""
    image_bytes: bytes | None = None


@dataclass
class SourceDocument:
    format: str
    pages: list[SourcePage] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    def full_text(self) -> str:
        blocks = []
        for page in self.pages:
            header = f"[PAGE {page.page_number}]"
            body = page.text.strip() or "(no extractable text)"
            blocks.append(f"{header}\n{body}")
        return "\n\n".join(blocks)


def detect_format(filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise UnsupportedAnswerKeyError(f"Unsupported answer-key type '{suffix}'. Supported: {supported}")
    return suffix.lstrip(".")


def convert_source(filename: str, data: bytes) -> SourceDocument:
    fmt = detect_format(filename)
    if fmt == "pdf":
        return _convert_pdf(data)
    if fmt == "docx":
        return _convert_docx(data)
    if fmt == "doc":
        return _convert_doc(data)
    return _convert_image(fmt, data)


# -- formats -----------------------------------------------------------------


def _convert_pdf(data: bytes) -> SourceDocument:
    import fitz  # PyMuPDF

    document = SourceDocument(format="pdf")
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:  # noqa: BLE001
        raise UnsupportedAnswerKeyError("PDF could not be opened") from exc
    zoom = 200 / 72  # render at ~200 DPI so scanned/image-only pages are readable by the VLM
    for index, page in enumerate(doc, start=1):
        image_bytes: bytes | None = None
        try:
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            image_bytes = pix.tobytes("png")
        except Exception:  # noqa: BLE001 - text-only keys must still parse
            logger.warning("PDF page %d could not be rendered to an image", index, exc_info=True)
        document.pages.append(SourcePage(page_number=index, text=page.get_text(), image_bytes=image_bytes))
    doc.close()
    return document


def _convert_docx(data: bytes) -> SourceDocument:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - guarded by extra
        raise UnsupportedAnswerKeyError("DOCX support not installed (pip install .[docparse])") from exc

    document = SourceDocument(format="docx")
    try:
        container = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise UnsupportedAnswerKeyError("DOCX could not be opened") from exc

    lines: list[str] = []
    for paragraph in container.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text)
    for table in container.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    document.pages.append(SourcePage(page_number=1, text="\n".join(lines)))
    return document


def _convert_doc(data: bytes) -> SourceDocument:
    """Legacy .doc requires LibreOffice; converts to DOCX then reuses that path."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        raise UnsupportedAnswerKeyError(
            "Legacy .doc files need LibreOffice installed on the server. Re-save the key as PDF or DOCX and upload again."
        )
    with tempfile.TemporaryDirectory(prefix="evalai-doc-") as tmp:
        src = Path(tmp) / "input.doc"
        src.write_bytes(data)
        try:
            completed = subprocess.run(  # noqa: S603
                [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, str(src)],
                capture_output=True,
                timeout=120,
            )
        except subprocess.TimeoutExpired as exc:
            raise UnsupportedAnswerKeyError("DOC conversion timed out") from exc
        converted = Path(tmp) / "input.docx"
        if completed.returncode != 0 or not converted.exists():
            raise UnsupportedAnswerKeyError("DOC conversion failed; re-save as PDF or DOCX")
        result = _convert_docx(converted.read_bytes())
        result.format = "doc"
        result.warnings.append("Converted from legacy .doc via LibreOffice")
        return result


def _convert_image(fmt: str, data: bytes) -> SourceDocument:
    try:
        image = Image.open(io.BytesIO(data))
        image.load()
    except Exception as exc:  # noqa: BLE001
        raise UnsupportedAnswerKeyError(f"{fmt.upper()} image could not be opened") from exc
    buffer = io.BytesIO()
    image.convert("RGB").save(buffer, format="PNG")
    return SourceDocument(
        format=fmt,
        pages=[SourcePage(page_number=1, text="", image_bytes=buffer.getvalue())],
        warnings=["Text will be read from the image by the AI parser"],
    )
